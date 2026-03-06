
from datetime import datetime
from pathlib import Path

from shared.logger import get_logger
from shared.errors import TraceLitError
from shared.utils.export_text import build_export_blocks, format_structured_text, inline_tokens_to_text

logger = get_logger(__name__)

def _confidence_rgb(confidence: str):
    from docx.shared import RGBColor

    if confidence == "high":
        return RGBColor(34, 139, 34)
    if confidence == "medium":
        return RGBColor(204, 153, 0)
    return RGBColor(204, 51, 51)

def _add_separator(doc, Pt, RGBColor) -> None:
    from docx.oxml.ns import qn

    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(2)
    sep.paragraph_format.space_after = Pt(2)
    pPr = sep._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single",
        qn("w:sz"): "4",
        qn("w:space"): "1",
        qn("w:color"): "D0D0D0",
    })
    pBdr.append(bottom)
    pPr.append(pBdr)

def _add_message_to_doc(doc, msg, Pt, RGBColor, WD_ALIGN_PARAGRAPH):
    role = msg.get("role", "user").upper()
    content = msg.get("content", "")
    havf_results = msg.get("havf_results") or []

    role_para = doc.add_paragraph()
    role_para.paragraph_format.space_before = Pt(6)
    role_para.paragraph_format.space_after = Pt(2)
    role_run = role_para.add_run(f"[{role}]")
    role_run.bold = True
    role_run.font.size = Pt(11)
    if role == "ASSISTANT":
        role_run.font.color.rgb = RGBColor(0, 51, 153)

    _render_blocks_to_doc(doc, build_export_blocks(content), Pt)

    if havf_results:
        _add_verification_section(doc, havf_results, Pt, RGBColor)

    _add_separator(doc, Pt, RGBColor)

def _add_verification_section(doc, havf_results, Pt, RGBColor):
    verif_para = doc.add_paragraph()
    verif_para.paragraph_format.space_before = Pt(4)
    verif_run = verif_para.add_run("Citation Verification:")
    verif_run.italic = True
    verif_run.font.size = Pt(9)
    verif_run.font.color.rgb = RGBColor(100, 100, 100)

    for r in havf_results:
        confidence = r.get("confidence", "low")
        raw_claim = r.get("claim", "")
        claim = raw_claim[:500].rsplit(" ", 1)[0] + "..." if len(raw_claim) > 500 else raw_claim
        paragraph_id = r.get("paragraph_id", "")
        chunk_type = r.get("chunk_type", "")
        score = r.get("score", 0)

        r_para = doc.add_paragraph()
        r_para.paragraph_format.space_before = Pt(1)
        r_para.paragraph_format.space_after = Pt(1)

        badge_run = r_para.add_run(f"  [{confidence.upper()}] ({score:.0%})")
        badge_run.bold = True
        badge_run.font.size = Pt(8)
        badge_run.font.color.rgb = _confidence_rgb(confidence)

        if paragraph_id:
            ref_run = r_para.add_run(f"  [{paragraph_id}]")
            ref_run.font.size = Pt(8)
            ref_run.font.color.rgb = RGBColor(100, 100, 100)

        if chunk_type and chunk_type != "text":
            type_run = r_para.add_run(f"  ({chunk_type})")
            type_run.font.size = Pt(8)
            type_run.font.color.rgb = RGBColor(100, 100, 100)

        claim_run = r_para.add_run(f'  "{format_structured_text(claim)}"')
        claim_run.font.size = Pt(8)
        claim_run.font.color.rgb = RGBColor(80, 80, 80)


def _set_paragraph_run_size(paragraph, size) -> None:
    for run in paragraph.runs:
        run.font.size = size


def _render_table_to_doc(doc, headers, rows, Pt) -> None:
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"

    for col_idx, header in enumerate(headers):
        cell = table.rows[0].cells[col_idx]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    for row_idx, row in enumerate(rows, start=1):
        for col_idx, value in enumerate(row[:len(headers)]):
            table.rows[row_idx].cells[col_idx].text = value
            for paragraph in table.rows[row_idx].cells[col_idx].paragraphs:
                _set_paragraph_run_size(paragraph, Pt(9))


def _render_blocks_to_doc(doc, blocks, Pt) -> None:
    for block in blocks:
        if block.kind == "heading":
            level = min(max(block.level, 1), 4)
            heading = doc.add_heading(block.text, level=level)
            heading.paragraph_format.space_after = Pt(4)
            _set_paragraph_run_size(heading, Pt(max(10, 15 - level)))
        elif block.kind == "bullet":
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.add_run(block.text)
            paragraph.paragraph_format.space_after = Pt(2)
            _set_paragraph_run_size(paragraph, Pt(10))
        elif block.kind == "table":
            _render_table_to_doc(doc, block.headers, block.rows, Pt)
            doc.add_paragraph()
        else:
            paragraph = doc.add_paragraph(inline_tokens_to_text(block.tokens))
            paragraph.paragraph_format.space_after = Pt(4)
            _set_paragraph_run_size(paragraph, Pt(10))


def _add_cited_media_section(doc, cited_assets, Pt) -> None:
    if not cited_assets:
        return

    try:
        from docx.shared import Inches
    except ImportError:
        Inches = None

    doc.add_heading("Cited Figures, Tables, and Formulas", level=2)
    for asset in cited_assets:
        label = f"[{asset.get('citation_id', '')}] {str(asset.get('chunk_type', '')).title()}"
        meta = f"{asset.get('paper_title', '')}"
        if asset.get("page_number"):
            meta += f" | page {asset.get('page_number')}"
        if asset.get("section_title"):
            meta += f" | {asset.get('section_title')}"

        heading = doc.add_paragraph()
        heading.add_run(label).bold = True
        if meta.strip():
            heading.add_run(f"  {meta}")
        _set_paragraph_run_size(heading, Pt(10))

        _render_blocks_to_doc(doc, build_export_blocks(str(asset.get("content", ""))), Pt)

        image_path = asset.get("image_path")
        if Inches and image_path and Path(image_path).exists():
            try:
                doc.add_picture(str(image_path), width=Inches(5.8))
            except Exception:
                pass
        doc.add_paragraph()

def _setup_doc_with_heading_and_date(doc, title, Pt, RGBColor, WD_ALIGN_PARAGRAPH):
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.paragraph_format.space_after = Pt(12)
    date_run = date_para.add_run(f"Exported on {datetime.now().strftime('%B %d, %Y')}")
    date_run.font.size = Pt(9)
    date_run.font.color.rgb = RGBColor(120, 120, 120)

def export_chat_to_docx(
    session_title: str,
    messages: list[dict],
    cited_assets: list[dict] | None,
    output_path: str | Path,
) -> Path:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as exc:
        raise TraceLitError(
            message="python-docx library not installed — Word export unavailable.",
            status_code=501,
        ) from exc

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    _setup_doc_with_heading_and_date(doc, session_title, Pt, RGBColor, WD_ALIGN_PARAGRAPH)

    for msg in messages:
        _add_message_to_doc(doc, msg, Pt, RGBColor, WD_ALIGN_PARAGRAPH)

    _add_cited_media_section(doc, cited_assets or [], Pt)

    try:
        doc.save(str(output_path))
    except Exception as exc:
        raise TraceLitError(
            message=f"Failed to generate Word document: {exc}",
            status_code=500,
        ) from exc

    logger.info(f"Exported chat DOCX to {output_path.name}")
    return output_path

def export_comparison_to_docx(
    title: str,
    comparison_content: str,
    paper_titles: list[str],
    comparison_table: list[dict] | None,
    cited_assets: list[dict] | None,
    output_path: str | Path,
) -> Path:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as exc:
        raise TraceLitError(
            message="python-docx library not installed — Word export unavailable.",
            status_code=501,
        ) from exc

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    _setup_doc_with_heading_and_date(doc, title, Pt, RGBColor, WD_ALIGN_PARAGRAPH)

    doc.add_heading("Papers Compared", level=2)
    for pt in paper_titles:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(pt)
        run.font.size = Pt(10)

    doc.add_heading("Comparison Analysis", level=2)
    if comparison_table:
        headers = ["Dimension", *paper_titles, "Synthesis"]
        rows = []
        for row in comparison_table:
            rows.append([
                str(row.get("dimension", "")),
                *[format_structured_text(str(cell.get("content", ""))) for cell in row.get("cells", [])],
                format_structured_text(str(row.get("synthesis", ""))),
            ])
        _render_table_to_doc(doc, headers, rows, Pt)
    else:
        _render_blocks_to_doc(doc, build_export_blocks(comparison_content), Pt)

    _add_cited_media_section(doc, cited_assets or [], Pt)

    try:
        doc.save(str(output_path))
    except Exception as exc:
        raise TraceLitError(
            message=f"Failed to generate comparison DOCX: {exc}",
            status_code=500,
        ) from exc

    logger.info(f"Exported comparison DOCX to {output_path.name}")
    return output_path

